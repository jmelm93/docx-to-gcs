import pathlib
import datetime
import markdown
import textwrap
from io import BytesIO
from google.cloud import storage
from google.cloud.storage import Blob
from htmldocx import HtmlToDocx
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
    
def set_cell_borders(cell, color="000000"):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders {}>'.format(nsdecls('w')) +
                          f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>' +
                          f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>' +
                          f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>' +
                          f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
                          '</w:tcBorders>')
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill, color="auto"):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}" w:color="{}"/>'.format(nsdecls('w'), fill, color))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def apply_paragraph_formatting(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(0), space_after=Pt(0)):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = space_before
    paragraph.paragraph_format.space_after = space_after


def apply_run_formatting(run, underline=False, color=RGBColor(0, 0, 0), font_size=None):
    run.font.underline = underline
    run.font.color.rgb = color
    if font_size is not None:
        run.font.size = font_size


def modify_docx_formatting(doc, final_path=None, save_local=True):
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Modify all tables
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            fill = "000000" if i == 0 else "FFFFFF"
            text_color = RGBColor(255, 255, 255) if i == 0 else RGBColor(0, 0, 0)
            font_size = Pt(13) if i == 0 else Pt(11)
            space_before, space_after = (Pt(4), Pt(4)) if i == 0 else (Pt(2), Pt(2))
            for cell in row.cells:
                set_cell_borders(cell)
                set_cell_shading(cell, fill=fill, color=text_color)
                for paragraph in cell.paragraphs:
                    apply_paragraph_formatting(paragraph, space_before=space_before, space_after=space_after)
                    for run in paragraph.runs:
                        apply_run_formatting(run, underline=(i == 0), color=text_color, font_size=font_size)

    # Adjust paragraph and run formatting
    header_sizes = {'Heading 1': Pt(24), 'Heading 2': Pt(21), 'Heading 3': Pt(18), 'Heading 4': Pt(15), 'Heading 5': Pt(12), 'Heading 6': Pt(12)}
    for paragraph in doc.paragraphs:
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        underline = paragraph.style.name == 'Heading 2'  
        space_before, space_after = (Pt(12), Pt(6)) if paragraph.style.name == 'Heading 2' else (Pt(4), Pt(0))
        apply_paragraph_formatting(paragraph, alignment=alignment, space_before=space_before, space_after=space_after)
        for run in paragraph.runs:
            font_size = header_sizes.get(paragraph.style.name, None)
            apply_run_formatting(run, color=RGBColor(0, 0, 0), font_size=font_size, underline=underline)
    
    if save_local and final_path:
        doc.save(final_path)



def html_to_docx(query, html_content):
    filename = query.replace(' ', '_') # convert 'query' to filename
    filename = ''.join(e for e in filename if e.isalnum()) # remove special characters
    full_filepath_no_extension = f'./output_files/{filename}'
    
    doc = Document()
    new_parser = HtmlToDocx()
    
    new_parser.add_html_to_document(html_content, doc)
    
    # reformat docx
    modify_docx_formatting(doc, f'{full_filepath_no_extension}.docx')


def html_to_docx_gcs(query, html_content, service_account_filename='service-account.json', bucket_name='content-maps'):
    # Construct the full path to the service account file based on the project root
    service_account_path = pathlib.Path(__file__).parent.parent / service_account_filename

    # Construct the filename
    filename = query.replace(' ', '_')  
    filename = ''.join(e for e in filename if e.isalnum()) 
    filename += f"_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}"
    filename += '.docx'
    
    # Create Document object
    doc = Document()
    new_parser = HtmlToDocx()
    new_parser.add_html_to_document(html_content, doc)

    # Reformat docx
    modify_docx_formatting(doc, save_local=False)

    # Save Document to a byte stream
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    # Google Cloud Storage client
    storage_client = storage.Client.from_service_account_json(service_account_path)
    bucket = storage_client.bucket(bucket_name)
    blob = Blob(filename, bucket)

    # Upload the file
    blob.upload_from_file(doc_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # Generate a public URL
    return blob.public_url


def markdown_to_docx_gcs(query, markdown_content, service_account_filename='service-account.json', bucket_name='content-maps'):
    """
    Converts Markdown content to a DOCX document and uploads it to Google Cloud Storage.

    This function first converts Markdown to HTML (with table support), then uses the same
    logic as `html_to_docx_gcs` to create and upload the DOCX file.
    """

    markdown_content = textwrap.dedent(markdown_content)

    # Convert Markdown to HTML with table support and other standard features
    html_content = markdown.markdown(
        markdown_content,
        extensions=["tables", "fenced_code"]
    )
    

    # Reuse the html_to_docx_gcs logic
    return html_to_docx_gcs(
        query=query,
        html_content=html_content,
        service_account_filename=service_account_filename,
        bucket_name=bucket_name
    )


if __name__ == "__main__":
    # Example usage
    query = "Example Query"
    markdown_content = """
    ## Competitive Keyword Targets:  

    | URL                                                         | Keywords                                                                 |
    |-------------------------------------------------------------|--------------------------------------------------------------------------|
    | https://www.geeksforgeeks.org/a-complete-guide-how-to-optimizetitle-tag/ | search engine optimized title (1600 searches/mo), title tag and seo (1000 searches/mo), title tag optimization (140 searches/mo), title tag optimisation (140 searches/mo), optimize title tags (140 searches/mo) |
    | https://www.gotchseo.com/title-tags/                        | title tag (3600 searches/mo), title tags (3600 searches/mo), title seo (1600 searches/mo), search engine optimization title (1600 searches/mo), seo title (1600 searches/mo) |
    | https://surferseo.com/blog/seo-title-tag/                   | what are seo (33100 searches/mo), title tags seo (1000 searches/mo), what is title in seo (720 searches/mo), title tag for seo (1000 searches/mo), title tags in seo (1000 searches/mo) |
    ## Recurring Topics:  

    **Title Tags and SEO Importance**  

    - Pages Mentioning Topic: [https://www.geeksforgeeks.org/a-complete-guide-how-to-optimizetitle-tag/], [https://surferseo.com/blog/seo-title-tag/]
    - Reason or Summary: Both pages emphasize the critical role of title tags in SEO, highlighting their impact on click-through rates, search engine ranking, and user experience.

    **Title Tag Best Practices**  

    - Pages Mentioning Topic: [https://www.geeksforgeeks.org/a-complete-guide-how-to-optimizetitle-tag/], [https://surferseo.com/blog/seo-title-tag/]
    - Reason or Summary: Both pages provide guidelines for optimizing title tags, such as keeping them concise, using relevant keywords, and ensuring uniqueness for each page.

    **Title Tag vs. Other Elements (H1, Meta Description)**  

    - Pages Mentioning Topic: [https://surferseo.com/blog/seo-title-tag/]
    - Reason or Summary: The distinction between title tags and other HTML elements like H1 and meta descriptions is discussed, emphasizing their different roles and the importance of alignment for SEO.

    ## Insights for SEO Optimization:  

    - **Focus on Title Tag Optimization**: Ensure that all web pages have unique, concise, and keyword-optimized title tags to improve search visibility and click-through rates.  
    - **Consistency Across Elements**: Align title tags with H1 tags and meta descriptions to provide a cohesive message to both users and search engines, enhancing relevance and user experience.  
    - **Regular Updates and Monitoring**: Periodically review and update title tags to reflect changes in content, keyword trends, and seasonal events, maintaining their effectiveness and relevance.
    """

    output = markdown_to_docx_gcs(query, markdown_content)
    print(output)